"""Word文件解析器 - 处理.docx文件"""
import os
import tempfile
import logging
from typing import List

from modules.parsers.base import FileParser
from modules.parsers import register_parser

logger = logging.getLogger("bidscrap")

@register_parser
class WordParser(FileParser):
    """Word文件解析器 - 支持.docx格式"""
    
    @classmethod
    async def parse(cls, file, column_index: int = 0, skip_rows: int = 1) -> List[str]:
        """解析Word文件提取企业名称"""
        companies = []
        
        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1])
        
        try:
            # 读取上传的文件内容
            content = await file.read()
            
            # 写入临时文件
            with open(temp_file.name, "wb") as f:
                f.write(content)
            
            # 使用python-docx处理Word文档
            import docx
            
            # 如果是.doc格式，无法直接处理，提示用户转换为.docx
            if os.path.splitext(file.filename)[1].lower() == '.doc':
                raise ValueError("暂不支持旧版Word(.doc)格式，请将文件另存为.docx格式后重试")
                
            doc = docx.Document(temp_file.name)
            
            # 获取文档中的所有段落文本
            all_text = []
            for i, para in enumerate(doc.paragraphs):
                if i >= skip_rows and para.text.strip():  # 跳过指定行数并忽略空行
                    all_text.append(para.text.strip())
            
            # 获取文档中的所有表格数据
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    if i >= skip_rows:  # 跳过表头
                        cells = [cell.text.strip() for cell in row.cells]
                        if len(cells) > column_index:
                            if cells[column_index]:
                                all_text.append(cells[column_index])
            
            # 清理企业名称
            companies = [str(text).strip() for text in all_text if str(text).strip()]
                
        except ImportError:
            logger.error("缺少python-docx库，无法解析Word文件")
            raise ValueError("系统未安装python-docx库，无法解析Word文件，请使用Excel或CSV格式")
        except Exception as e:
            logger.error(f"解析Word文件时出错: {str(e)}")
            raise e
        finally:
            # 删除临时文件
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
        
        return companies
    
    @classmethod
    def can_handle(cls, file_extension: str) -> bool:
        """判断是否可以处理该文件类型"""
        return file_extension.lower() in ['.docx']
    
    @classmethod
    @property
    def name(cls) -> str:
        """解析器名称"""
        return "word" 